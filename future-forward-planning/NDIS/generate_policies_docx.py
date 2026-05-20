"""
Generate Verification Module Policies — Word Document (.docx)
SC Partnering Pty Ltd (trading as Future Forward Planning)
Policies: PM-POL-002, PM-POL-003, PM-POL-004, PM-POL-005
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = '1B2A4A'
BLUE   = '2E6DA4'
LGRAY  = 'F0F4FA'
ALTROW = 'EEF2F8'
MGRAY  = 'CCCCCC'
DGRAY  = '333333'
WHITE  = 'FFFFFF'
NOTEBG = 'EEF4FB'


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_table_borders(table, color=MGRAY):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_para_border(paragraph, hex_color=BLUE, bg=NOTEBG):
    """Add a left border and background shading to a paragraph (note style)."""
    pPr = paragraph._p.get_or_add_pPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)

    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_run(paragraph, text, bold=False, italic=False,
            size=10, color=DGRAY, font='Calibri'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


# ── Inline markdown parser → paragraph runs ───────────────────────────────────

def add_inline_text(paragraph, text, base_size=10, base_color=DGRAY, base_italic=False):
    """Parse **bold** and *italic* markdown and add as runs."""
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+?\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            add_run(paragraph, part[2:-2], bold=True, size=base_size, color=base_color)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            add_run(paragraph, part[1:-1], italic=True, size=base_size, color=base_color)
        else:
            add_run(paragraph, part, italic=base_italic, size=base_size, color=base_color)


# ── Document setup ─────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Modify built-in styles so they match the design
    styles = doc.styles

    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after  = Pt(8)

    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(4)

    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(3)

    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(DGRAY)
    normal.paragraph_format.space_before = Pt(3)
    normal.paragraph_format.space_after  = Pt(3)

    return doc


# ── Header and footer ──────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(hp, 'SC Partnering Pty Ltd  |  Future Forward Planning',
            bold=True, size=8, color='888888')
    add_run(hp, '        Verification Module Policies  |  v1.0  |  May 2026',
            size=8, color='888888')

    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, 'ABN 24 697 821 408     |     ', size=8, color='888888')

    # Page number field
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run = fp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string('888888')
    run._r.append(fld_begin)
    run._r.append(instrText)
    run._r.append(fld_end)

    add_run(fp, '     |     CONFIDENTIAL — NDIS Submission Document', size=8, color='888888')


# ── Cover page ─────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Main title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(6)
    add_run(t, 'Verification Module Policies', bold=True, size=28, color=NAVY)

    # Subtitle
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(4)
    add_run(s, 'SC Partnering Pty Ltd\n', size=13, color=DGRAY)
    add_run(s, 'trading as Future Forward Planning', bold=True, size=13, color=DGRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Policy index table
    rows = [
        ('Policy', 'Title', 'Practice Standard'),
        ('PM-POL-001', 'Incident Management Policy', '1.9'),
        ('PM-POL-002', 'Complaints and Feedback Management Policy', '1.10'),
        ('PM-POL-003', 'Risk Management Policy', '1.7'),
        ('PM-POL-004', 'Human Resources and Worker Screening Policy', '1.8'),
    ]
    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(tbl)
    col_widths = [Cm(3.2), Cm(9.5), Cm(3.3)]
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = col_widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            if r_idx == 0:
                set_cell_bg(cell, NAVY)
                add_run(p, text, bold=True, size=10, color=WHITE)
            else:
                bg = LGRAY if r_idx % 2 == 1 else ALTROW
                set_cell_bg(cell, bg)
                add_run(p, text, size=10, color=DGRAY)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Details table
    details = [
        ('Business:',       'SC Partnering Pty Ltd (trading as Future Forward Planning)'),
        ('ABN:',            '24 697 821 408'),
        ('Version:',        '1.0 — Initial Release'),
        ('Effective Date:', '15 May 2026'),
        ('Review Date:',    '15 May 2027'),
        ('Approved By:',    'Saran Likitkunawong — Director'),
        ('Classification:', 'Confidential — NDIS Submission Document'),
    ]
    dtbl = doc.add_table(rows=len(details), cols=2)
    set_table_borders(dtbl, MGRAY)
    for r_idx, (label, value) in enumerate(details):
        row = dtbl.rows[r_idx]
        row.cells[0].width = Cm(3.8)
        row.cells[1].width = Cm(12.2)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(3)
        add_run(row.cells[0].paragraphs[0], label, bold=True, size=10, color=DGRAY)
        add_run(row.cells[1].paragraphs[0], value,            size=10, color=DGRAY)

    doc.add_page_break()


# ── Markdown table → docx table ───────────────────────────────────────────────

def add_md_table(doc, lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\|[-:| ]+\|$', stripped):
            continue
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    set_table_borders(tbl)

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx in range(col_count):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            text = row_data[c_idx] if c_idx < len(row_data) else ''
            if r_idx == 0:
                set_cell_bg(cell, NAVY)
                add_run(p, text, bold=True, size=9, color=WHITE)
            else:
                bg = WHITE if r_idx % 2 == 1 else LGRAY
                set_cell_bg(cell, bg)
                add_inline_text(p, text, base_size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Section content parser ─────────────────────────────────────────────────────

META_PREFIXES = (
    '**Policy Number:**', '**Effective Date:**', '**Review Date:**',
    '**Owner:**', '**Version Control:**',
)


def parse_section(doc, md_text):
    lines = md_text.split('\n')
    i = 0
    pending_bullets = []
    in_table = False
    table_buf = []

    def flush_bullets():
        for b in pending_bullets:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline_text(p, b, base_size=10)
        pending_bullets.clear()

    def flush_table():
        nonlocal in_table
        if table_buf:
            add_md_table(doc, table_buf)
        table_buf.clear()
        in_table = False

    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()
        i  += 1

        # Table
        if s.startswith('|'):
            flush_bullets()
            in_table = True
            table_buf.append(s)
            continue
        if in_table:
            flush_table()

        # Skip blank / nbsp / separators
        if not s or s == '&nbsp;' or s == '---':
            flush_bullets()
            continue

        # H1 policy title
        if s.startswith('# POLICY'):
            flush_bullets()
            doc.add_heading(s.lstrip('# '), level=1)
            continue

        # H2
        if s.startswith('## '):
            flush_bullets()
            doc.add_heading(s[3:], level=2)
            continue

        # H3
        if s.startswith('### '):
            flush_bullets()
            doc.add_heading(s[4:], level=3)
            continue

        # H4
        if s.startswith('#### '):
            flush_bullets()
            p = doc.add_paragraph()
            add_run(p, s[5:], bold=True, size=10, color=DGRAY)
            continue

        # Metadata lines
        if any(s.startswith(pfx) for pfx in META_PREFIXES):
            flush_bullets()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline_text(p, s, base_size=10)
            continue

        # Blockquote / note
        if s.startswith('> '):
            flush_bullets()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            set_para_border(p)
            add_inline_text(p, s[2:], base_size=9.5, base_color='2c4a6e', base_italic=True)
            continue

        # Checkbox bullet
        if re.match(r'^- \[.?\]', s):
            checked = bool(re.match(r'^- \[[xX]\]', s))
            text = re.sub(r'^- \[.?\]\s*', '', s)
            mark = '☑' if checked else '☐'
            pending_bullets.append(f'{mark} {text}')
            continue

        # Bullet
        if re.match(r'^[-*]\s', s):
            pending_bullets.append(s[2:])
            continue

        # Numbered list
        if re.match(r'^\d+\.\s', s):
            flush_bullets()
            text = re.sub(r'^\d+\.\s+', '', s)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_inline_text(p, text, base_size=10)
            continue

        # Regular paragraph
        flush_bullets()
        p = doc.add_paragraph()
        add_inline_text(p, s, base_size=10)

    flush_bullets()
    if in_table:
        flush_table()


# ── Extract policies 2–5 ───────────────────────────────────────────────────────

def extract_policies(filepath):
    with open(filepath, encoding='utf-8') as f:
        content = f.read()

    policies = {}
    for n in [1, 2, 3, 4]:
        start = content.find(f'# POLICY {n}:')
        end   = content.find(f'# POLICY {n + 1}:')
        if start == -1:
            continue
        chunk = content[start: end if end != -1 else len(content)]
        chunk = re.sub(r'\n---\n---.*$', '', chunk.rstrip(), flags=re.DOTALL)
        policies[n] = chunk

    return policies


# ── Main ──────────────────────────────────────────────────────────────────────

def generate(output_path, suite_path):
    doc = setup_document()
    add_header_footer(doc)
    add_cover(doc)

    policies = extract_policies(suite_path)
    for n in [1, 2, 3, 4]:
        if n not in policies:
            print(f'WARNING: Policy {n} not found')
            continue
        parse_section(doc, policies[n])
        if n < 4:
            doc.add_page_break()

    doc.save(output_path)
    print(f'Done: {output_path}')


if __name__ == '__main__':
    base = r'C:\Users\Saran Likitkunawong\VScode\NDIS'
    generate(
        output_path=os.path.join(base, 'Verification Module Policies.docx'),
        suite_path=os.path.join(base, 'Policy Suite - Draft Templates.md'),
    )
